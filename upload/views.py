from django.shortcuts import render
from django.http import HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
import docx
import pdfplumber
import numpy as np
import pandas as pd
import re
import string
import os
import openpyxl
import nltk
import textblob
from textblob import TextBlob
from spellchecker import SpellChecker
nltk.download('wordnet')
from nltk.stem import WordNetLemmatizer
wnl = WordNetLemmatizer()
from nltk.corpus import stopwords
nltk.download('stopwords')
from nltk.tokenize import word_tokenize

# Create your views here.
@csrf_exempt
def index(request):
    result="Upload service up"
    if request.method == 'POST':
        print(request.FILES)
        result = my_program(request.FILES)
        return JsonResponse(result, safe=False)
    return HttpResponse(result)
    # return render(request, "index.html")
    
    
def extract_text(path):
    
    if path.endswith('.docx'):
        doc = docx.Document(path)
        text = ""
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
            text = '\n'.join(fullText)
        
        return text
                
    elif path.endswith('.pdf'):
        text = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                fullText = page.extract_text()
                text += '\n' + fullText
                
        return text

# make everything lowercase to reduce vocabulary

def text_lowercase(text):
    return text.lower()

#remove punctuations

def remove_punctuation(text):
    translator = str.maketrans('', '', string.punctuation)
    return text.translate(translator)


# remove whitespace from text

def remove_whitespace(text):
    return " ".join(text.split())

# remove stopwords function
def remove_stopwords(text):
    stop_words = set(stopwords.words("english"))
    word_tokens = word_tokenize(text)
    filtered_text = [word for word in word_tokens if word not in stop_words]
    return filtered_text


def preprocessing(text):
    
    text_lower = text_lowercase(text)
    text_nonpunc = remove_punctuation(text_lower)
    text_nospace = remove_whitespace(text_nonpunc)
    text_nospace = text_nospace.replace('\n',' ')
    text_cleaned = remove_stopwords(text_nospace)
    text_cleaned = " ".join(text_cleaned)
    text_cleaned = re.compile('<.*?>').sub('', text_cleaned) 
    text_cleaned = re.compile('[%s]' % re.escape(string.punctuation)).sub(' ', text_cleaned)  
    text_cleaned = re.sub('\s+', ' ', text_cleaned)  
    text_cleaned = re.sub(r'\[[0-9]*\]',' ',text_cleaned) 
    text_cleaned = re.sub(r'[^\w\s]', '', str(text_cleaned).lower().strip())
    text_cleaned = re.sub(r'\d',' ',text_cleaned) 
    text_cleaned = re.sub(r'\s+',' ',text_cleaned)
    
    return text_cleaned

def lemmatizer(keyword):
    finallist = []
    for keywords in keyword:
        finallist.append(wnl.lemmatize(keywords))
        
    return finallist

df = pd.DataFrame()

excel_path = "RPA Data.xlsx"
data = pd.read_excel(excel_path, sheet_name = None)
keyword_data = data['Keyword']
master_keyword_list = keyword_data['Keyword'].tolist()
designation_data = data['Designation']
developer_data = data['Developer']
developer_data = developer_data.fillna("")


def my_program(files):
    fileList = []
    for f in files.keys():
        newFile = {}
        
        if files[f].name.endswith('.docx'):
            doc = docx.Document(files[f])
            #doc = docx2txt.process(files[f])
            text = ""
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
                text = '\n'.join(fullText)
        
            newFile['fileName'] = files[f].name
            newFile['content'] = text
            fileList.append(newFile)
                
        elif files[f].name.endswith('.pdf'):
            text = ""
            fullText = []
            with pdfplumber.open(files[f]) as pdf:
                for page in pdf.pages:
                    fullText = page.extract_text()
                    text += '\n' + fullText
                
            newFile['fileName'] = files[f].name
            newFile['content'] = text
            fileList.append(newFile)
            
        else:
            
            newFile['fileName'] = files[f].name
            newFile['content'] = files[f].read().decode("utf-8")
            fileList.append(newFile)
                  
    #print(x)
   
    #extracting text from pdf/docx files
    #jd_text = extract_text(jd_path)
    jd_text = fileList[0]['content']
    
    #word count for jd
    wordCount_JD = len(jd_text.split())
    
    resume_text = fileList[1]['content']
            
    #regex to find whether text contains email address
    pattern = '\w[\w\.-]*@\w[\w\.-]+\.\w+'
    emails = re.findall(pattern, resume_text)

    if len(emails)>0:
        bonus_email='True'

    else:
        bonus_email='False'

    #regex to find whether text contains mobile number
    pattern = "(?:\+ *)?\d[\d\- ]{10,}\d"
    mobile = re.findall(pattern, resume_text)

    if len(mobile)>0:
        bonus_mobile='True'

    else:
        bonus_mobile='False'

    #regex to find whether text contains linkedIn address
    pattern = "ttps:\/\/www.linkedin.com\/in\/"
    linkedIn = re.findall(pattern, resume_text)

    if len(linkedIn)>0:
        bonus_linkedIn='True'

    else:
        bonus_linkedIn='False'

    #word count for resume
    wordCount_Resume = len(resume_text.split())

    #performing pre processing on both file data

    jd_cleaned = preprocessing(jd_text)
    resume_cleaned = preprocessing(resume_text)

    #matching hard skills, soft skills, certifications
    skills = []

    def extractor(var):
        for i in range(0,len(developer_data[var])):
            if developer_data[var][i] != "":
                if resume_text.find(developer_data[var][i]) >= 0:
                    skills.append(developer_data[var][i])

    hard_skills = []
    soft_skills = []
    certifications = []

    extractor('Soft Skills')
    soft_skills = skills
    skills = []

    extractor('Hard Skills')
    hard_skills = skills
    skills = []

    extractor('Certifications')
    certifications = skills
    skills = []

    #extracting keywords from both text and creating final list
    resume_keyword = resume_cleaned.split()
    jd_keyword = jd_cleaned.split()

    #lemmatizing words in list
    resume_list = lemmatizer(resume_keyword)
    jd_list = lemmatizer(jd_keyword)

    resume_finalList = pd.unique(resume_list).tolist()
    jd_finalList = pd.unique(jd_list).tolist()

    #creation of required list
    req_list = []

    def comp(list1, list2):
        for val in list1:
            if val in list2:
                req_list.append(val)

    #this is the list of requirements in JD
    comp(jd_finalList,master_keyword_list)

    #creating final available list in resume for that JD
    final_list = []

    def comp_2(list1, list2):
        for val in list1:
            if val in list2:
                final_list.append(val)

    comp_2(resume_finalList,req_list)

    #calculating score acquired by resume for this JD

    score = (len(final_list)/len(req_list))*100

    df = df.append({"JD Filepath" : path, "Resume Filepath" : var, 'Score' : score, "Word Count Resume" : wordCount_Resume,
                    "Word Count JD" : wordCount_JD, "Contains Mobile" : bonus_mobile, "Contains Email" : bonus_email,
                    "Contains LinkedIn" : bonus_linkedIn, "Soft Skills" : soft_skills, "Hard Skills" : hard_skills,
                   "Certifications" : certifications}, ignore_index = True)

    return score
